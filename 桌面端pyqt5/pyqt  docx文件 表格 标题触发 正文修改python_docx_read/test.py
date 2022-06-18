#读写标题，读写表格
import docx
from docx import Document

obj = Document("1.docx")

# for ii in range(len(obj.paragraphs)):
for p in obj.paragraphs:
    # p = obj.paragraphs[ii]
    style_name = p.style.name
    if style_name.startswith('Heading'):
        print(style_name,p.text,sep=':')
        #删除
        p.clear()
        #添加
        p.add_run("added")


# 遍历所有表格
tables = obj.tables #获取文件中的表格集
print(len(tables))
for table in tables[:]:
    print("-"*20)
    for i, row in enumerate(table.rows[:]):   # 读每行
        row_content = []
        for cell in row.cells[:]:  # 读一行中的所有单元格
            c = cell.text
            #修改数据
            # cell.text = "231"
            row_content.append(c)
        print (row_content) #以列表形式导出每一行数据

    print(table.cell(0,0).text)

obj.save("2.docx")

pass